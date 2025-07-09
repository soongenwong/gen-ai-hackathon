# backend.py

import os
import sys
import io
import logging
import pandas as pd
from typing import Dict, Any, Optional
import mimetypes
import tempfile

# Additional imports for file types
try:
    import pdfplumber
except ImportError:
    pdfplumber = None
try:
    import docx
except ImportError:
    docx = None
try:
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup
except ImportError:
    epub = None
    ebooklib = None
    BeautifulSoup = None

# ──────────────────────────────────────────────────────────────────────────────
# SDV imports - comprehensive functionality
# ──────────────────────────────────────────────────────────────────────────────
from sdv.metadata import SingleTableMetadata
from sdv.single_table.ctgan import CTGANSynthesizer
from sdv.single_table.copulagan import CopulaGANSynthesizer
from sdv.single_table.ctgan import TVAESynthesizer
from sdv.single_table.copulas import GaussianCopulaSynthesizer
from sdv.evaluation.single_table import evaluate_quality, run_diagnostic
from sdv.sampling import Condition

# ──────────────────────────────────────────────────────────────────────────────
# Logging setup
# ──────────────────────────────────────────────────────────────────────────────
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)
handler = logging.StreamHandler()
formatter = logging.Formatter(
    "[%(asctime)s] %(levelname)s in %(module)s: %(message)s"
)
handler.setFormatter(formatter)
logger.addHandler(handler)


# ──────────────────────────────────────────────────────────────────────────────
# Core Functions
# ──────────────────────────────────────────────────────────────────────────────
def load_dataframe(uploaded_file) -> pd.DataFrame:
    """
    Load the uploaded file (CSV, Excel, PDF, DOCX, EPUB, TXT) into a pandas DataFrame and validate it's not empty.
    Throws clear ValueErrors for unsupported column types and malformed files.
    """
    try:
        # Try to get file name and extension
        file_name = getattr(uploaded_file, 'name', None)
        file_ext = os.path.splitext(file_name)[-1].lower() if file_name else ''
        mime_type = mimetypes.guess_type(file_name)[0] if file_name else None

        # Read file into DataFrame based on extension or mime type
        if file_ext in ['.csv', '.tsv'] or (mime_type and 'csv' in mime_type):
            sep = '\t' if file_ext == '.tsv' else ','
            df = pd.read_csv(uploaded_file, sep=sep)
        elif file_ext in ['.xls', '.xlsx'] or (mime_type and 'excel' in mime_type):
            df = pd.read_excel(uploaded_file)
        elif file_ext == '.txt' or (mime_type and 'text' in mime_type):
            # Try to read as CSV, else as plain text
            try:
                df = pd.read_csv(uploaded_file)
            except Exception:
                uploaded_file.seek(0)
                text = uploaded_file.read().decode('utf-8', errors='ignore') if hasattr(uploaded_file, 'read') else uploaded_file.getvalue().decode('utf-8', errors='ignore')
                df = pd.DataFrame({'text': text.splitlines()})
        elif file_ext == '.pdf':
            if pdfplumber is None:
                raise ImportError("pdfplumber is required for PDF support. Please install it.")
            uploaded_file.seek(0)
            with pdfplumber.open(uploaded_file) as pdf:
                tables = []
                for page in pdf.pages:
                    page_tables = page.extract_tables()
                    for table in page_tables:
                        tables.append(pd.DataFrame(table[1:], columns=table[0]))
                if tables:
                    df = pd.concat(tables, ignore_index=True)
                else:
                    # Fallback: extract text
                    text = '\n'.join(page.extract_text() or '' for page in pdf.pages)
                    df = pd.DataFrame({'text': text.splitlines()})
        elif file_ext == '.docx':
            if docx is None:
                raise ImportError("python-docx is required for DOCX support. Please install it.")
            doc = docx.Document(uploaded_file)
            tables = []
            for table in doc.tables:
                data = []
                for row in table.rows:
                    data.append([cell.text for cell in row.cells])
                tables.append(pd.DataFrame(data))
            if tables:
                df = pd.concat(tables, ignore_index=True)
            else:
                # Fallback: extract all text
                text = '\n'.join([p.text for p in doc.paragraphs])
                df = pd.DataFrame({'text': text.splitlines()})
        elif file_ext == '.epub':
            if epub is None or ebooklib is None or BeautifulSoup is None:
                raise ImportError("ebooklib and beautifulsoup4 are required for EPUB support. Please install them.")
            # Save to a temporary file
            with tempfile.NamedTemporaryFile(delete=True, suffix='.epub') as tmp:
                tmp.write(uploaded_file.read())
                tmp.flush()
                book = epub.read_epub(tmp.name)
            text = ''
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    text += soup.get_text() + '\n'
            df = pd.DataFrame({'text': text.splitlines()})
        else:
            raise ValueError(f"Unsupported file type: {file_ext or 'unknown'}. Supported: CSV, Excel, PDF, DOCX, EPUB, TXT.")

        # Check if DataFrame is empty
        if df.empty:
            raise ValueError("Uploaded file is empty - no data rows found.")
        
        # Check if DataFrame has no columns
        if len(df.columns) == 0:
            raise ValueError("Uploaded file has no columns.")
        
        # Check for completely unnamed columns
        unnamed_cols = [col for col in df.columns if str(col).startswith('Unnamed:')]
        if len(unnamed_cols) == len(df.columns):
            raise ValueError("File appears to have no proper column headers - all columns are unnamed.")
        
        # Validate column types - check for unsupported data types
        unsupported_types = []
        supported_dtypes = ['int64', 'float64', 'object', 'bool', 'datetime64[ns]', 'category']
        
        for col in df.columns:
            dtype_str = str(df[col].dtype)
            # Check for complex data types that SDV can't handle
            if any(unsupported in dtype_str.lower() for unsupported in ['complex', 'timedelta']):
                unsupported_types.append(f"{col} ({dtype_str})")
            
            # Check for object columns that might contain mixed types
            if df[col].dtype == 'object':
                # Sample a few values to check for problematic types
                sample_values = df[col].dropna().head(10)
                for val in sample_values:
                    if isinstance(val, (list, dict, tuple, set)):
                        unsupported_types.append(f"{col} (contains {type(val).__name__} objects)")
                        break
        
        if unsupported_types:
            raise ValueError(
                f"Unsupported column types detected that cannot be processed by SDV: "
                f"{', '.join(unsupported_types)}. "
                f"Supported types are: numeric (int, float), text (string), boolean, and datetime."
            )
        
        # Check for columns with all null values
        null_columns = df.columns[df.isnull().all()].tolist()
        if null_columns:
            raise ValueError(
                f"Columns with all null/missing values detected: {', '.join(null_columns)}. "
                f"Please remove these columns or provide data for them."
            )
        
        # Check minimum data requirements for synthetic data generation
        if len(df) < 2:
            raise ValueError(
                f"Insufficient data for synthetic generation - found only {len(df)} row(s). "
                f"At least 2 rows are required."
            )
        
        # Warn about small datasets (but don't fail)
        if len(df) < 10:
            logger.warning(
                f"Small dataset detected ({len(df)} rows). "
                f"Synthetic data quality may be limited. Consider providing more data for better results."
            )
        
        logger.info(f"Loaded DataFrame with {len(df)} rows, {len(df.columns)} cols.")
        return df
        
    except pd.errors.EmptyDataError:
        raise ValueError("File is empty or contains no data.")
    
    except pd.errors.ParserError as e:
        raise ValueError(f"Malformed file - parsing failed: {str(e)}")
    
    except UnicodeDecodeError as e:
        raise ValueError(
            f"File encoding issue - unable to read file: {str(e)}. "
            f"Please ensure the file is saved with UTF-8 encoding."
        )
    
    except FileNotFoundError:
        raise ValueError("File not found. Please check the file path.")
    
    except PermissionError:
        raise ValueError("Permission denied - unable to read the file.")
    
    except Exception as e:
        # Catch any other unexpected errors and provide a clear message
        logger.exception("Failed to load DataFrame")
        raise ValueError(f"Failed to load file: {str(e)}")


def create_metadata_with_constraints(df: pd.DataFrame, constraints_config: Optional[Dict] = None) -> SingleTableMetadata:
    """
    Create comprehensive metadata with optional constraints.
    """
    metadata = SingleTableMetadata()
    metadata.detect_from_dataframe(df)
    
    constraints_count = 0  # Initialize the counter
    
    # Add constraints if provided
    if constraints_config:
        try:
            # Import constraint classes when needed
            from sdv.constraints.tabular import Range, Positive
            
            if 'range_constraints' in constraints_config:
                for col, (min_val, max_val) in constraints_config['range_constraints'].items():
                    if col in df.columns:
                        # Fix: Use correct parameter names for Range constraint
                        constraint = Range(
                            column=col,
                            low=min_val,
                            high=max_val
                        )
                        metadata.add_constraint(constraint)
                        constraints_count += 1
            
            if 'positive_constraints' in constraints_config:
                for col in constraints_config['positive_constraints']:
                    if col in df.columns:
                        # Fix: Use correct parameter name for Positive constraint
                        constraint = Positive(column=col)
                        metadata.add_constraint(constraint)
                        constraints_count += 1
        except ImportError:
            logger.warning("Constraint classes not available, skipping constraints")
        except Exception as e:
            logger.warning(f"Error adding constraints: {e}, skipping constraints")
    
    logger.info(f"Metadata created with {constraints_count} constraints")
    return metadata


def get_synthesizer(model_type: str, metadata: SingleTableMetadata, **kwargs):
    """
    Factory function to create different types of synthesizers.
    """
    synthesizers = {
        'ctgan': CTGANSynthesizer,
        'copulagan': CopulaGANSynthesizer,
        'tvae': TVAESynthesizer,
        'gaussiancopula': GaussianCopulaSynthesizer
    }
    
    if model_type not in synthesizers:
        raise ValueError(f"Unknown model type: {model_type}. Available: {list(synthesizers.keys())}")
    
    synthesizer_class = synthesizers[model_type]
    return synthesizer_class(metadata=metadata, **kwargs)


def generate_synthetic_data(
    df: pd.DataFrame,
    n_rows: int,
    model_type: str = 'ctgan',
    privacy_level: float = 0.5,
    constraints_config: Optional[Dict] = None,
    conditional_sampling: Optional[Dict] = None
) -> pd.DataFrame:
    """
    Generate synthetic data using various SDV synthesizers with advanced features.
    
    Args:
        df: Input DataFrame
        n_rows: Number of synthetic rows to generate
        model_type: Type of synthesizer ('ctgan', 'copulagan', 'tvae', 'gaussiancopula')
        privacy_level: 0.0 (full fidelity) → 1.0 (maximum privacy/faster run)
        constraints_config: Dictionary with constraint specifications
        conditional_sampling: Dictionary for conditional sampling
    """
    try:
        # 1) Build metadata with constraints
        metadata = create_metadata_with_constraints(df, constraints_config)
        logger.info("SDV metadata detected with constraints.")

        # 2) Configure model parameters based on privacy level
        model_params = {}
        if model_type in ['ctgan', 'copulagan', 'tvae']:
            base_epochs = 300 if model_type == 'ctgan' else 100
            epochs = max(1, int(base_epochs * (1 - privacy_level)))
            model_params['epochs'] = epochs
            logger.info(f"Training {model_type.upper()} for {epochs} epochs (privacy={privacy_level}).")

        # 3) Create and fit synthesizer
        synthesizer = get_synthesizer(model_type, metadata, **model_params)
        synthesizer.fit(df)
        
        # 4) Generate synthetic data with optional conditional sampling
        if conditional_sampling:
            conditions = []
            for col, value in conditional_sampling.items():
                if col in df.columns:
                    condition = Condition(
                        num_rows=n_rows // len(conditional_sampling),
                        column_values={col: value}
                    )
                    conditions.append(condition)
            
            if conditions:
                try:
                    synthetic_df = synthesizer.sample_conditions(conditions)
                    logger.info(f"Generated {len(synthetic_df)} synthetic rows with conditions.")
                except AttributeError:
                    # Fallback to regular sampling if conditional sampling not available
                    synthetic_df = synthesizer.sample(num_rows=n_rows)
                    logger.info(f"Generated {len(synthetic_df)} synthetic rows (conditional sampling not available).")
            else:
                synthetic_df = synthesizer.sample(num_rows=n_rows)
                logger.info(f"Generated {len(synthetic_df)} synthetic rows.")
        else:
            synthetic_df = synthesizer.sample(num_rows=n_rows)
            logger.info(f"Generated {len(synthetic_df)} synthetic rows.")
        
        return synthetic_df

    except Exception:
        logger.exception("Synthetic generation failed")
        raise RuntimeError("Could not generate synthetic data")


def evaluate_synthetic_quality(real_data: pd.DataFrame, synthetic_data: pd.DataFrame) -> Dict[str, Any]:
    """
    Evaluate the quality of synthetic data using SDV metrics.
    """
    try:
        # Create metadata for evaluation
        metadata = SingleTableMetadata()
        metadata.detect_from_dataframe(real_data)
        
        # Run quality evaluation
        quality_report = evaluate_quality(real_data, synthetic_data, metadata)
        
        # Run diagnostic evaluation
        diagnostic_report = run_diagnostic(real_data, synthetic_data, metadata)
        
        logger.info("Quality evaluation completed.")
        
        # Fix: Handle the get_details() method properly
        quality_details = {}
        try:
            # Try to get available property names
            if hasattr(quality_report, 'get_properties'):
                properties = quality_report.get_properties()
                for prop in properties:
                    try:
                        quality_details[prop] = quality_report.get_details(prop)
                    except:
                        continue
        except:
            quality_details = {"error": "Could not retrieve quality details"}
        
        diagnostic_results = {}
        try:
            if hasattr(diagnostic_report, 'get_properties'):
                properties = diagnostic_report.get_properties()
                for prop in properties:
                    try:
                        diagnostic_results[prop] = diagnostic_report.get_details(prop)
                    except:
                        continue
        except:
            diagnostic_results = {"error": "Could not retrieve diagnostic results"}
        
        return {
            'quality_score': quality_report.get_score(),
            'quality_details': quality_details,
            'diagnostic_results': diagnostic_results,
            'diagnostic_score': diagnostic_report.get_score()
        }
    
    except Exception:
        logger.exception("Quality evaluation failed")
        return {"error": "Quality evaluation failed"}


def generate_privacy_report(synthesizer, real_data: pd.DataFrame, synthetic_data: pd.DataFrame) -> Dict[str, Any]:
    """
    Generate a privacy analysis report.
    """
    try:
        # Basic privacy metrics
        privacy_metrics = {
            'data_size_ratio': len(synthetic_data) / len(real_data),
            'column_coverage': len(set(synthetic_data.columns)) / len(set(real_data.columns)),
            'unique_values_preservation': {}
        }
        
        # Check unique value preservation for categorical columns
        for col in real_data.select_dtypes(include=['object']).columns:
            if col in synthetic_data.columns:
                real_unique = set(real_data[col].unique())
                synthetic_unique = set(synthetic_data[col].unique())
                preservation_ratio = len(real_unique.intersection(synthetic_unique)) / len(real_unique)
                privacy_metrics['unique_values_preservation'][col] = preservation_ratio
        
        logger.info("Privacy report generated.")
        return privacy_metrics
    
    except Exception:
        logger.exception("Privacy report generation failed")
        return {"error": "Privacy report generation failed"}


def dataframe_to_csv_bytes(df: pd.DataFrame) -> bytes:
    """
    Convert DataFrame to CSV bytes for Streamlit download.
    """
    buf = io.StringIO()
    df.to_csv(buf, index=False)
    data = buf.getvalue().encode("utf-8")
    logger.info("Converted DataFrame to CSV bytes.")
    return data


# ──────────────────────────────────────────────────────────────────────────────
# Optional CLI for local testing
# ──────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import sys as _sys

    if len(_sys.argv) < 3:
        print("Usage: python main.py <input_csv> <num_rows> [model_type] [privacy_level]")
        _sys.exit(1)

    input_csv = _sys.argv[1]
    num_rows = int(_sys.argv[2])
    model_type = _sys.argv[3] if len(_sys.argv) > 3 else 'ctgan'
    privacy_level = float(_sys.argv[4]) if len(_sys.argv) > 4 else 0.5

    try:
        df = load_dataframe(input_csv)
        synth = generate_synthetic_data(df, num_rows, model_type, privacy_level)
        print(f"\nSynthetic data generated using {model_type.upper()}:")
        print(synth.head())
        
        # Evaluate quality
        quality_results = evaluate_synthetic_quality(df, synth)
        print(f"\nQuality Score: {quality_results.get('quality_score', 'N/A')}")
        
    except Exception as e:
        print(f"Error: {e}")
        _sys.exit(1)